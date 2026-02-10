from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def create_pdf(resume_data):
    """Generate ATS-friendly PDF resume"""
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=0.75*inch, leftMargin=0.75*inch,
                           topMargin=0.75*inch, bottomMargin=0.75*inch)
    
    # Container for PDF elements
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#1E88E5'),
        spaceAfter=6,
        alignment=TA_CENTER
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#333333'),
        spaceAfter=6,
        spaceBefore=12,
        borderWidth=1,
        borderColor=colors.HexColor('#1E88E5'),
        borderPadding=5
    )
    
    # Name (Title)
    name = Paragraph(resume_data.get('name', 'Your Name'), title_style)
    elements.append(name)
    
    # Contact Info
    contact_info = f"{resume_data.get('phone', '')} | {resume_data.get('email', '')}"
    contact = Paragraph(contact_info, styles['Normal'])
    elements.append(contact)
    elements.append(Spacer(1, 0.2*inch))
    
    # Professional Summary
    if resume_data.get('summary'):
        elements.append(Paragraph('PROFESSIONAL SUMMARY', heading_style))
        elements.append(Paragraph(resume_data['summary'], styles['Normal']))
        elements.append(Spacer(1, 0.1*inch))
    
    # Skills
    if resume_data.get('skills'):
        elements.append(Paragraph('SKILLS', heading_style))
        elements.append(Paragraph(resume_data['skills'], styles['Normal']))
        elements.append(Spacer(1, 0.1*inch))
    
    # Experience
    if resume_data.get('experience'):
        elements.append(Paragraph('EXPERIENCE', heading_style))
        exp_text = resume_data['experience'].replace('\n', '<br/>')
        elements.append(Paragraph(exp_text, styles['Normal']))
        elements.append(Spacer(1, 0.1*inch))
    
    # Projects
    if resume_data.get('projects'):
        elements.append(Paragraph('PROJECTS', heading_style))
        proj_text = resume_data['projects'].replace('\n', '<br/>')
        elements.append(Paragraph(proj_text, styles['Normal']))
        elements.append(Spacer(1, 0.1*inch))
    
    # Education
    if resume_data.get('education'):
        elements.append(Paragraph('EDUCATION', heading_style))
        edu_text = resume_data['education'].replace('\n', '<br/>')
        elements.append(Paragraph(edu_text, styles['Normal']))
        elements.append(Spacer(1, 0.1*inch))
    
    # Certifications
    if resume_data.get('certifications'):
        elements.append(Paragraph('CERTIFICATIONS & ACHIEVEMENTS', heading_style))
        cert_text = resume_data['certifications'].replace('\n', '<br/>')
        elements.append(Paragraph(cert_text, styles['Normal']))
    
    # Build PDF
    doc.build(elements)
    
    buffer.seek(0)
    return buffer.getvalue()

def create_docx(resume_data):
    """Generate ATS-friendly DOCX resume"""
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Name (Title)
    name = doc.add_heading(resume_data.get('name', 'Your Name'), 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.runs[0].font.color.rgb = RGBColor(30, 136, 229)
    
    # Contact Info
    contact = doc.add_paragraph(f"{resume_data.get('phone', '')} | {resume_data.get('email', '')}")
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Professional Summary
    if resume_data.get('summary'):
        doc.add_heading('PROFESSIONAL SUMMARY', 1)
        doc.add_paragraph(resume_data['summary'])
    
    # Skills
    if resume_data.get('skills'):
        doc.add_heading('SKILLS', 1)
        doc.add_paragraph(resume_data['skills'])
    
    # Experience
    if resume_data.get('experience'):
        doc.add_heading('EXPERIENCE', 1)
        doc.add_paragraph(resume_data['experience'])
    
    # Projects
    if resume_data.get('projects'):
        doc.add_heading('PROJECTS', 1)
        doc.add_paragraph(resume_data['projects'])
    
    # Education
    if resume_data.get('education'):
        doc.add_heading('EDUCATION', 1)
        doc.add_paragraph(resume_data['education'])
    
    # Certifications
    if resume_data.get('certifications'):
        doc.add_heading('CERTIFICATIONS & ACHIEVEMENTS', 1)
        doc.add_paragraph(resume_data['certifications'])
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()